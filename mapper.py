import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

RIGHT_DOOR_HALLS = [] #place the numbers of halls with a right door alignment here separated by commas.
                      #the mapper will accordingly queue students from the door according to government map requests.


                    #grades partaking in your exams, the numbers must be aligned with the mentioned year from 7-12
ARABIC_GRADES = {
    7: "اول متوسط", 8: "ثاني متوسط", 9: "ثالث متوسط",
    10: "رابع اعدادي", 11: "خامس اعدادي", 12: "سادس"
}




#__________________________-do not edit below here-__________________________
def get_names_from_txt(file_name):
    if not os.path.exists(file_name): return []
    with open(file_name, 'r', encoding='utf-8') as f:
        content = f.read()
        return [n.strip() for n in content.split(',') if n.strip()]

def set_landscape(section):
    """Sets the page orientation to Landscape."""
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

def fill_hall_header(doc, hall_num):
    for section in doc.sections:
        set_landscape(section) # Force landscape on individual docs
        header = section.header
        for paragraph in header.paragraphs:
            if "رقم القــــــــــــــــاعة" in paragraph.text:
                paragraph.text = paragraph.text.replace("(    )", f"({hall_num})")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def format_cell_large(cell, s12_name, y_student):
    """Sets text to Name/Class with LARGER font."""
    line1 = f"{s12_name} / {ARABIC_GRADES[12]}"
    line2 = f"{y_student['name']} / {y_student['grade']}" if y_student else ""
    
    cell.text = f"{line1}\n{line2}"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(14)  # Increased font size
            run.font.bold = True    # Bold for better visibility

def create_seating_landscape():
    y12 = get_names_from_txt('year_12.txt')
    younger = []
    for y in range(7, 12):
        names = get_names_from_txt(f'year_{y}.txt')
        for n in names:
            younger.append({"name": n, "grade": ARABIC_GRADES[y]})

    master_doc = Document()

    for hall_num in range(1, 15):
        start_idx = (hall_num - 1) * 21
        hall_12s = y12[start_idx : start_idx + 21]
        if not hall_12s: break

        doc = Document('base.docx')
        fill_hall_header(doc, hall_num)
        
        table = doc.tables[0]
        table.width = Inches(9.0) # Stretch table for landscape
        
        col_order = [2, 1, 0] if hall_num in RIGHT_DOOR_HALLS else [0, 1, 2]

        idx = 0
        for col in col_order:
            for row in range(7):
                if idx < len(hall_12s):
                    s12 = hall_12s[idx]
                    y_idx = (hall_num - 1) * 21 + idx
                    y_stu = younger[y_idx] if y_idx < len(younger) else None
                    format_cell_large(table.cell(row, col), s12, y_stu)
                    idx += 1

        doc.save(f'Seating_Hall_{hall_num}.docx')

        # MASTER DOC Logic
        if hall_num == 1:
            set_landscape(master_doc.sections[0])
        else:
            new_sec = master_doc.add_section()
            set_landscape(new_sec)
            new_sec.header.is_linked_to_previous = False

        # Clone header and table to Master
        curr_sec = master_doc.sections[-1]
        for i, p in enumerate(doc.sections[0].header.paragraphs):
            while len(curr_sec.header.paragraphs) <= i: curr_sec.header.add_paragraph()
            curr_sec.header.paragraphs[i].text = p.text
            curr_sec.header.paragraphs[i].alignment = p.alignment

        master_table = master_doc.add_table(rows=7, cols=3)
        master_table.style = 'Table Grid'
        master_table.width = Inches(9.0)
        
        idx = 0
        for col in col_order:
            for row in range(7):
                if idx < len(hall_12s):
                    s12 = hall_12s[idx]
                    y_idx = (hall_num - 1) * 21 + idx
                    y_stu = younger[y_idx] if y_idx < len(younger) else None
                    format_cell_large(master_table.cell(row, col), s12, y_stu)
                    idx += 1

    master_doc.save('All_Halls_Landscape_Master.docx')
    print("Files generated in Landscape with large fonts.")

if __name__ == "__main__":
    create_seating_landscape()